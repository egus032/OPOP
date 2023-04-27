import docx
import glob
import os
import pandas as pd


g_path = glob.glob(r'C:\Users\User\Downloads\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А\*.docx')

# g_path = glob.glob(r'C:\Users\User\Downloads\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А\*.docx')

# g_path = os.walk(r'C:\Users\User\Downloads\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А')

g_list = []

g_dict = {}

keys = None

counter = 0

for g_file in g_path:
    g_file_short = g_file.replace(r'C:\Users\User\Downloads\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А' + '\\', '')
    if g_file_short.find('~$') == -1:
        g_doc = docx.Document(g_file)
        g_table = g_doc.tables
        g_dict[g_file_short] = g_table
        g_dict.update
        for key in g_dict:
            for i, t in enumerate(g_dict[key]):
                if i >= 0:
                    print(key, i, t)
                    my_list = [['' for i in range(len(t.columns))] for j in range(len(t.rows))]
                    for i, row in enumerate(t.rows):
                        for j, cell in enumerate(row.cells):
                            # if cell.text:
                                # print(cell.text)
                            my_list[i][j] = cell.text
                    g_list.append(my_list)
                    # print(g_list)
                    df = pd.DataFrame(g_list)
                    print(df)
                    df.to_excel('лист.xlsx')
                    # for row, col in zip(t.rows, t.columns):
                    #     # for j, cell in enumerate(row.cells):
                    #     # for r, c in zip(row.cells, col.cells):
                    #     #     print(r.text, c.text)
                    #     #     break
                    #     # print(row.cells, col.cells)
                    #     for r, c in zip(row.cells, col.cells):
                    #         for p_1, p_2 in zip(r.paragraphs, c.paragraphs):
                    #             print(p_1.text, p_2.text)
                    #             break

                    

# print(g_dict)

# for key in g_dict:
#     for i, t in enumerate(g_dict[key]):
#         if i >= 0:
#             print(key, i, t)
        
                   
            
            
            # if i != 0:
            #     keys = list(text)
            #     print(keys)
            #     continue
            # row_data = dict(zip(keys, text))
            # row_data.update
            # g_list.append(row_data)
    # print(key, g_list)
    # break
        # df = [['' for i in range(len(t.columns))] for j in range(len(t.rows))]
        # for i, row in enumerate(t.rows):
        #     for j, cell in enumerate(row.cells):
        #         if cell.text:
        #             # print(cell.text)
        #             df[i][j] = cell.text
        #             print(df)
        # data = [[cell.text for cell in row.cells] for row in t.rows]
        # print(data)
# df = pd.DataFrame(data)
            
