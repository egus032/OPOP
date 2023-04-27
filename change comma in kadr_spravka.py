from docx import Document
import glob
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

g_path = r'C:\Users\User\Desktop\Рабочая папка\от Геращенковой\2023\апрель 2023\Кадровые справки в ворд_2_rename и запятые\*'

g_glob = glob.glob(g_path)

tables = []

error_list = []

for g_file in g_glob:
    document = Document(g_file)
    for t, table in enumerate(document.tables):
        if t == 1:
            for c, col in enumerate(table.columns):
                if c == 2:
                    for cell in col.cells:
                        if '.' in cell.text:
                            cell.text = cell.text.replace('.', ',')
                            for paragraph in cell.paragraphs:
                                style = document.styles['Normal']
                                cell.style = style
                                font = style.font
                                font.name = 'Times New Roman'
                                font.size = Pt(11)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print(cell.text)
    document.save(g_file)
                
        