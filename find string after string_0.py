import docx
import glob
import os
import pandas as pd

g_fold = r'C:\Users\User\Desktop\Рабочая папка\от Геращенковой\2023\февраль 2023\РПД_2018_10.05.04_аид_О'

g_path = glob.glob(g_fold + '\\*.*')

g_list = []

g_dict = {}

keys = None

counter = 0

para_list = []

f_word = '(далее – дисциплина) ориентирована на формирование у обучающихся компетенций в'

for g_file in g_path:
    if g_file.find('~$Д_') == -1:
        g_file_short = g_file.replace(g_fold + '\\', '')
        g_doc = docx.Document(g_file)
        if g_file.find(g_file_short) != -1:
            g_dict[g_file_short] = list(filter(None, [paragraph.text for paragraph in g_doc.paragraphs]))
            g_dict.update()
            
for id_x, text_in in g_dict.items():
    for t in text_in:
        if f_word in t:
            counter += 1
            g_string = f'{id_x}; {t}'
            g_list.append(g_string)
            print(counter, g_string)

df = pd.DataFrame(data=g_list)
print(df)
df.to_excel('код РПД_Дисциплина.xlsx')

df_1 = pd.DataFrame()


           
# l = [dict(zip([g_file_short],[paragraph.text])) for paragraph in g_doc.paragraphs]

            # for t, table in enumerate(g_table):
            #     # print(g_file_short, t, table)
            #     if t == 2:
            #         df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            #         # print(df)
            #         for r, row in enumerate(table.rows):
            #             for c, cell in enumerate(row.cells):
            #                 if cell.text:
            #                     df[r][c] = cell.text
            #                 df_tables.append(pd.DataFrame(df))
            #             print(df_tables)
            #         break


        
                    
    
# for key in get_dict_with_tables(g_path):
#     for i, t in enumerate(g_dict[key]):
#         if i >= 10 and i <= 15:
#             header_row = t.rows[0].cells
#             for j, h in enumerate(header_row):
#                 print(j, h.text)
                # if j == 0 or j <= 1:
                #     print(j, h.text)
                    # if j == 0:
                    #     g_dict_0[j] = h.text
                    #     g_dict_0.update
                        
                        
                    # if j == 1:
                    #     g_dict_1[j] = h.text
                    #     g_dict_1.update
                        
                    # df_0 = pd.DataFrame.from_dict(g_dict_0, orient='index')
                    # df_1 = pd.DataFrame.from_dict(g_dict_1, orient='index')
                    # print(df_0, df_1)

                    
                    
                        
            

        # # print(i, t.row_cells(row_idx=1))
        # for r, c in zip(t.row_cells(1), t.column_cells(1)):
        #     # print(key, i, t, j, c.text)
        #     print(r.text, c.text)
            

        # for key in g_dict:
        #     for i, t in enumerate(g_dict[key]):
        #         if i >= 0:
        #             print(key, i, t)
        #             my_list = [['' for i in range(len(t.columns))] for j in range(len(t.rows))]
        #             for i, row in enumerate(t.rows):
        #                 for j, cell in enumerate(row.cells):
        #                     # if cell.text:
        #                         # print(cell.text)
        #                     my_list[i][j] = cell.text
        #             g_list.append(my_list)
        #             # print(g_list)
        #             df = pd.DataFrame(g_list)
        #             print(df)
        #             df.to_excel('лист.xlsx')
                    # for row, col in zip(t.rows, t.columns):
                    #     # for j, cell in enumerate(row.cells):
                    #     # for r, c in zip(row.cells, col.cells):
                    #     #     print(r.text, c.text)
                    #     #     break
                    #     # print(row.cells, col.cells)
                    #     for r, c in zip(row.cells, col.cells):
                    #         for p_1, p_2 in zip(r.paragraphs, c.paragraphs):
                    #             print(p_1.text, p_2.text)
                    #             break

                    

# print(g_dict)

# for key in g_dict:
#     for i, t in enumerate(g_dict[key]):
#         if i >= 0:
#             print(key, i, t)
        
                   
            
            
            # if i != 0:
            #     keys = list(text)
            #     print(keys)
            #     continue
            # row_data = dict(zip(keys, text))
            # row_data.update
            # g_list.append(row_data)
    # print(key, g_list)
    # break
        # df = [['' for i in range(len(t.columns))] for j in range(len(t.rows))]
        # for i, row in enumerate(t.rows):
        #     for j, cell in enumerate(row.cells):
        #         if cell.text:
        #             # print(cell.text)
        #             df[i][j] = cell.text
        #             print(df)
        # data = [[cell.text for cell in row.cells] for row in t.rows]
        # print(data)
# df = pd.DataFrame(data)
            
