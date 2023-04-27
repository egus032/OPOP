

import docx
from docx.shared import Pt
import glob

word_folder = glob.glob(r'D:\OPOP\3++_РП_ГИА_магистр_теги\*.docx')

list_paragraphs = []

for g_file in word_folder:
    doc = docx.Document(g_file)
    for d in doc.paragraphs:
        print(d)
    break

    # # doc.paragraphs[1].style.font.name = 'Times New Roman'
    # # doc.paragraphs[1].style.font.size = 20
    # substring = g_file.replace('D:\\OPOP\\ОПОПы_3+_бакалавр\\', ' ')
    # doc.save(f'new_{substring}')
    # # print(doc.paragraphs[2].runs[1].text)
    
    
        
    
    
    
        
