import docx
import glob
import os
import pandas as pd

g_fold = r'C:\Users\User\Desktop\Рабочая папка\от Геращенковой\2023\февраль 2023\РПД_2018_10.05.04_аид_О'

# g_fold = r'C:\Users\User\Downloads\РПД_2021_22.04.01_ирм_О_Пилюшина Г.А'

g_path = glob.glob(g_fold + '\\*.*')

g_list = []

g_dict = {}

keys = None

counter = 0

df_tables = []

first_list = ['Наименование раздела (темы) дисциплины', 'Трудоемкость, час.', 'Трудоемкость, час.', 'Трудоемкость, час.', 'Трудоемкость, час.', 'Трудоемкость, час.']
second_list = []

for g_file in g_path:
    g_file_short = g_file.replace(g_fold + '\\', '')
    if g_file_short.find('~$') == -1 or g_file_short.startswith('!_'):
        g_doc = docx.Document(g_file)
        for t, table in enumerate(g_doc.tables):
            for table_text in table.row_cells(0):
                if t == 10:
                    print(g_file_short, t, table_text.text)
#                     g_dict.setdefault(g_file_short, []).append(table_text.text)
#                     g_dict.update()
#         print(g_file_short)
#         print(g_dict)

# df = pd.DataFrame.from_dict(g_dict, orient='index')
# print(df)
# # df.to_excel('Обработка РПД 01.03.23_1.xlsx')



            
#             # if t >= 10 and t < 11:
#             for first_row, second_row in zip(table.row_cells(0), table.row_cells(1)):
                
#                 g_dict.setdefault(g_file_short, []).append(first_row.text)
#                 g_dict.update()
#         print(g_dict)
#         for key in g_dict.keys():
#             for f_l in first_list:
#                 if f_l in g_dict[key]:
#                     True
                    
#             print(first_list)
#         break
            

# # for g_file in g_path:
# #         g_file_short = g_file.replace(g_fold + '\\', '')
# #         if g_file_short.find('~$') == -1:
# #             g_doc = docx.Document(g_file)
# #             g_table = g_doc.tables
# #             for table in g_table:
# #                 g_dict.setdefault(g_file_short, []).append(table)
# #                 g_dict.update()
           
# #             for i, key in enumerate(g_dict.keys()):
# #                 g_string = f'{i, key, len(g_dict[key])}'
            
# #             print(g_string)
            
            