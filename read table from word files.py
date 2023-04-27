
from itertools import tee
import pandas as pd
from docx import Document
import glob

g_folder = glob.glob(r'C:\Users\User\Desktop\ВПО\*\*\*\*.docx')

df_list = []

for g_file in g_folder:
    g_doc = Document(g_file)
    g_result = [p.text for p in g_doc.paragraphs if p.text != '']
    
    for table in g_doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                data = [[cell.text for cell in row.cells] for row in table.rows]
                
        df = pd.DataFrame(data)
        df['Файл'] = g_file
        df['Тип сотрудника'] = g_result[1]
    df_list.append(df)
    print(df_list)
df_1 = pd.concat(df_list)
df_1.to_excel('2019-2020 из ворда.xlsx')
                
                